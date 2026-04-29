"""
Vantag Documentation Package Generator — shared library.
Provides consistent styling, cover pages, TOCs, footers across all 20+ docs.
"""
from __future__ import annotations
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand ─────────────────────────────────────────────────────────────────
BRAND_NAME = "Vantag Retail Intelligence Platform"
BRAND_SHORT = "Vantag"
VERSION = "1.0"
DATE_STR = datetime.utcnow().strftime("%B %Y")
DOCS_ROOT = os.path.dirname(os.path.abspath(__file__))

# Colors
C_PRIMARY = RGBColor(0x5B, 0x21, 0xB6)
C_ACCENT = RGBColor(0x10, 0xB9, 0x81)
C_DARK = RGBColor(0x1F, 0x29, 0x37)
C_MUTED = RGBColor(0x6B, 0x72, 0x80)
C_CODE_BG = RGBColor(0xF3, 0xF4, 0xF6)

# ── Helpers ───────────────────────────────────────────────────────────────
def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D1D5DB")
        borders.append(b)
    tbl_pr.append(borders)


def _add_footer(doc, page_footer):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(page_footer)
    run.font.size = Pt(8)
    run.font.color.rgb = C_MUTED


def _apply_base_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def new_doc(title, subtitle, doc_code, category):
    doc = Document()
    _apply_base_style(doc)
    _add_footer(doc, f"{BRAND_SHORT} — {doc_code} — Confidential  |  Page ")

    # Cover page
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(BRAND_NAME.upper())
    r.font.size = Pt(14)
    r.font.color.rgb = C_PRIMARY
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n" + title)
    r.font.size = Pt(32)
    r.font.color.rgb = C_DARK
    r.bold = True

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.font.size = Pt(14)
        r.font.color.rgb = C_MUTED
        r.italic = True

    for _ in range(6):
        doc.add_paragraph()

    # Metadata table
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_borders(meta)
    rows = [
        ("Document Code", doc_code),
        ("Category", category),
        ("Version", VERSION),
        ("Date", DATE_STR),
        ("Status", "Approved for Release"),
    ]
    for i, (k, v) in enumerate(rows):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = v
        _shade_cell(meta.rows[i].cells[0], "F3F4F6")
        for c in meta.rows[i].cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    for c in meta.rows[0].cells:
        c.width = Inches(2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("© 2026 Vantag. Confidential. Not for redistribution.")
    r.font.size = Pt(9)
    r.font.color.rgb = C_MUTED
    r.italic = True

    doc.add_page_break()
    return doc


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.color.rgb = C_PRIMARY
    r.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.color.rgb = C_DARK
    r.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.color.rgb = C_PRIMARY
    r.bold = True
    p.paragraph_format.space_before = Pt(6)


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.line_spacing = 1.25
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for r in p.runs:
        r.font.size = Pt(11)
    return p


def code_block(doc, text, language=""):
    # Mermaid / code rendered as monospaced block
    if language:
        cap = doc.add_paragraph()
        r = cap.add_run(f"Diagram ({language}):")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = C_MUTED

    table = doc.add_table(rows=1, cols=1)
    _set_borders(table)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, "F3F4F6")
    for line in text.strip("\n").split("\n"):
        p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = C_DARK


def table_rows(doc, header, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    _set_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = h
        _shade_cell(cell, "5B21B6")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(widths):
                t.rows[ri].cells[ci].width = Inches(w)
    return t


def toc(doc, items):
    h2(doc, "Table of Contents")
    t = doc.add_table(rows=len(items), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (sec, title) in enumerate(items):
        t.rows[i].cells[0].text = sec
        t.rows[i].cells[1].text = title
        t.rows[i].cells[0].width = Inches(0.8)
        for c in t.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    doc.add_page_break()


def save(doc, folder, filename):
    out = os.path.join(DOCS_ROOT, folder, filename)
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    print(f"  Saved: {folder}/{filename}")
    return out
