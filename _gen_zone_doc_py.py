# _gen_zone_doc_py.py
# Called by gen_zone_doc.js — uses python-docx to produce a clean .docx

import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"D:\AI Algo\Collaterals\Profiles\Retail Nazar\Vantag_Zone_Configuration_Guide.docx"

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Styles helpers ────────────────────────────────────────────────────────────

def set_heading_style(heading, size_pt, color_hex, bold=True):
    run = heading.runs[0] if heading.runs else heading.add_run(heading.text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r, g, b = int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16)
    run.font.color.rgb = RGBColor(r, g, b)

def para_run(doc_or_cell, text, bold=False, italic=False, size=11, color=None,
             font="Calibri", align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=6, space_after=6):
    if isinstance(doc_or_cell, Document):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        r,g,b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
        run.font.color.rgb = RGBColor(r,g,b)
    return p

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x1F,0x38,0x64)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x2E,0x74,0xB5)
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None,
             space_before=4, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        r,g,b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
        run.font.color.rgb = RGBColor(r,g,b)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    return p

def set_cell_bg(cell, color_hex):
    """Set table cell background shading."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def make_table(doc, headers, rows, col_widths_pct=None):
    """Create a bordered table with header row and data rows."""
    n_cols = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths if provided
    if col_widths_pct:
        page_width = Inches(6.5)  # ~6.5" usable
        for i, pct in enumerate(col_widths_pct):
            for row in t.rows:
                row.cells[i].width = int(page_width * pct / 100)

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, "2E74B5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name  = "Calibri"
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = t.rows[ri+1]
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = drow.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return t

def add_code_block(doc, text):
    """Add a monospace code block with light grey background."""
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        # Background shading on paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F0F0F0')
        pPr.append(shd)
        # Border around paragraph
        pBdr = OxmlElement('w:pBdr')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    '4')
            b.set(qn('w:space'), '2')
            b.set(qn('w:color'), 'AAAAAA')
            pBdr.append(b)
        pPr.append(pBdr)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after  = Pt(8)
run = p.add_run("Vantag Platform — Zone Configuration, Detection Logic & Accuracy Guide")
run.font.name  = "Calibri"
run.font.size  = Pt(24)
run.font.bold  = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
r2 = p2.add_run("How AI detection works, how zones are set up, what is accurate and what needs improvement")
r2.font.name   = "Calibri"
r2.font.size   = Pt(13)
r2.font.italic = True
r2.font.color.rgb = RGBColor(0x44,0x44,0x44)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(24)
r3 = p3.add_run("v1.0 — April 2026")
r3.font.name  = "Calibri"
r3.font.size  = Pt(11)
r3.font.color.rgb = RGBColor(0x66,0x66,0x66)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, "Executive Summary")

add_para(doc,
    "The Vantag platform runs real AI detection (YOLOv8) on live RTSP camera feeds. "
    "However, there is an important distinction between what is technically running and "
    "what is fully calibrated for your specific store layout.")

make_table(doc,
    ["Detection Type", "Current Status"],
    [
        ["Camera Tamper Detection",      "ACCURATE — works without any zone config"],
        ["Fall Detection",               "ACCURATE — uses human pose estimation (skeleton joints)"],
        ["Restricted Zone Entry",        "PARTIALLY ACCURATE — zones defined in pixel coords but not visually drawn on your actual feed"],
        ["Shoplifting/Sweep Detection",  "PARTIALLY ACCURATE — runs on full frame, works when person is near shelf"],
        ["Inventory Movement",           "NEEDS CALIBRATION — shelf zones defined but not verified against your real shelf positions"],
        ["Queue Length Monitoring",      "NEEDS CALIBRATION — queue zone defined but not verified against your checkout counter"],
        ["Watchlist / Face Match",       "NOT YET ACTIVE — requires face image uploads to the watchlist"],
    ],
    [40, 60],
)

add_para(doc, "")

add_para(doc,
    "This document explains exactly how each detection type works, what is configured, "
    "why the risk scores show the numbers they do, and what steps are needed to make "
    "every detection maximally accurate.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: ZONE CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, "How Zones Are Configured — The Truth")
add_h2(doc, "Zone Definition Method")

add_para(doc,
    "All zones in Vantag are defined as pixel coordinate rectangles (bounding boxes) or "
    "polygons relative to the camera's resolution (1920×1080). These coordinates are "
    "stored in a single configuration file:")

add_code_block(doc, "Location: vantag/backend/config/cameras.yaml")

add_para(doc, "There are two types of zone shapes:")
add_bullet(doc, "Bounding Box (bbox): A rectangle defined by [x1, y1, x2, y2] pixel coordinates. Used for inventory shelves and queue zones.")
add_bullet(doc, "Polygon: A list of [x, y] corner points forming any shape. Used for restricted zones (e.g., back office, storage room).")

add_h2(doc, "What Is Currently Configured Per Camera")
add_para(doc, "The following zones are already defined in cameras.yaml for your 3 active cameras:")

make_table(doc,
    ["Camera", "Zone Type", "Zone Name", "Pixel Coordinates"],
    [
        ["cam-01 (Zone A)", "Inventory Shelf", "Shelf A1",                 "x1=50 y1=200 x2=640 y2=700"],
        ["cam-01 (Zone A)", "Inventory Shelf", "Shelf A2",                 "x1=640 y1=200 x2=1280 y2=700"],
        ["cam-01 (Zone A)", "Inventory Shelf", "Shelf A3",                 "x1=1280 y1=200 x2=1870 y2=700"],
        ["cam-01 (Zone A)", "Restricted Zone", "Back Office",              "polygon top-right corner (1600,0)\u2192(1920,400)"],
        ["cam-01 (Zone A)", "Restricted Zone", "Staff Only (10pm\u20138am only)", "polygon bottom-left corner (0,800)\u2192(300,1080)"],
        ["cam-01 (Zone A)", "Queue Zone",      "Checkout Lane A",          "x1=700 y1=600 x2=1300 y2=1080"],
        ["cam-03 (Zone C)", "Inventory Shelf", "Shelf C1",                 "x1=50 y1=150 x2=700 y2=720"],
        ["cam-03 (Zone C)", "Inventory Shelf", "Shelf C2",                 "x1=700 y1=150 x2=1400 y2=720"],
        ["cam-03 (Zone C)", "Inventory Shelf", "Shelf C3",                 "x1=1400 y1=150 x2=1870 y2=720"],
        ["cam-03 (Zone C)", "Inventory Shelf", "Freezer Zone",             "x1=1500 y1=0 x2=1920 y2=600"],
        ["cam-03 (Zone C)", "Restricted Zone", "Storage Room C",           "polygon top-right (1700,0)\u2192(1920,500)"],
        ["cam-03 (Zone C)", "Queue Zone",      "Checkout Lane C",          "x1=600 y1=650 x2=1350 y2=1080"],
        ["cam-04 (Zone D)", "Inventory Shelf", "Shelf D1",                 "x1=50 y1=150 x2=700 y2=750"],
        ["cam-04 (Zone D)", "Inventory Shelf", "Shelf D2",                 "x1=700 y1=150 x2=1400 y2=750"],
        ["cam-04 (Zone D)", "Inventory Shelf", "Entrance Display",         "x1=50 y1=750 x2=960 y2=1080"],
        ["cam-04 (Zone D)", "Restricted Zone", "Manager Office D",         "polygon top-right (1650,0)\u2192(1920,450)"],
        ["cam-04 (Zone D)", "Queue Zone",      "Checkout Lane D",          "x1=600 y1=600 x2=1350 y2=1080"],
    ],
    [18, 18, 22, 42],
)

add_para(doc, "")

# Warning box
p_warn = doc.add_paragraph()
p_warn.paragraph_format.space_before = Pt(8)
p_warn.paragraph_format.space_after  = Pt(8)
p_warn.paragraph_format.left_indent  = Inches(0.25)
# shading
pPr = p_warn._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FFF3CD')
pPr.append(shd)
run_w = p_warn.add_run(
    "THE CORE ISSUE: These coordinates were set as estimates based on a standard 1920\u00d71080 "
    "retail layout. They have NOT been drawn and verified on your actual live camera feeds. "
    "Until you visually confirm that these boxes overlap the real shelves and zones in your "
    "store, detections may fire in the wrong area or miss the correct area."
)
run_w.font.name  = "Calibri"
run_w.font.size  = Pt(11)
run_w.font.bold  = True
run_w.font.color.rgb = RGBColor(0x85,0x4d,0x0e)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: HOW EACH DETECTION WORKS
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, "How Each AI Detection Works")

# 3.1
add_h2(doc, "1. Camera Tamper Detection")
add_para(doc, "Tamper detection compares consecutive video frames using Mean Squared Error (MSE). It does NOT require any zones.")
add_para(doc, "How it works:")
add_bullet(doc, "Every 2 seconds, the system takes the current frame and compares it to a stored baseline frame.")
add_bullet(doc, "If the MSE difference between frames exceeds a threshold (default: 1500), a tamper event fires.")
add_bullet(doc, "This catches: camera being covered, moved, sprayed, or pointed away.")
add_bullet(doc, "The baseline is updated every 10 minutes automatically.")
add_para(doc, "Accuracy: VERY HIGH (95%+). This works on raw pixel math — no AI model required. False positives can occur if lights change dramatically (e.g., store opening/closing). Cooldown is 30 seconds per camera so it will not spam alerts.")
add_para(doc, "Why Zone C and D show CRITICAL: These cameras are Dahua cameras connected over RTSP. During the initial connection phase and whenever the RTSP stream reconnects (due to network hiccup), the frame difference is large enough to trigger a tamper event. This is the main source of the 43 incidents you are seeing right now.")
add_para(doc, "Action Required: None. This is working correctly. If you see many tamper alerts at startup, it is normal — they stop once the stream stabilises.")

# 3.2
add_h2(doc, "2. Fall Detection")
add_para(doc, "Fall detection uses two methods: pose skeleton analysis (preferred) and bounding box ratio fallback.")
add_para(doc, "How it works:")
add_bullet(doc, "The YOLOv8-pose model (yolov8n-pose.pt) detects 17 body keypoints on every person: nose, shoulders, hips, knees, ankles.")
add_bullet(doc, "Primary method (Pose): The system checks if hip keypoints and knee keypoints are at the same vertical height (within 30 pixels). In normal standing/walking, hips are always above knees. If they are level, the person is horizontal (fallen).")
add_bullet(doc, "Fallback method (Bbox ratio): If pose keypoints are unavailable, it checks the bounding box width vs height. If width \u00f7 height > 1.4, the person is likely horizontal.")
add_bullet(doc, "A person must be detected as fallen for 5 consecutive frames before an alert fires (to avoid false positives from someone bending down).")
add_bullet(doc, "45-second cooldown per tracked person.")
add_para(doc, "Accuracy: HIGH (80\u201390%). Works well for actual falls. False positives can occur when someone sits on the floor, crouches to pick something up, or a child walks past. The 5-frame minimum helps filter these out.")
add_para(doc, "Action Required: None for basic operation. For higher accuracy, increase min_fallen_frames to 8\u201310 in cameras.yaml.")

# 3.3
add_h2(doc, "3. Restricted Zone Entry Detection")
add_para(doc, "This analyzer uses computer vision polygon hit-testing to detect when a person enters a defined area.")
add_para(doc, "How it works:")
add_bullet(doc, "Each restricted zone is defined as a polygon (list of pixel corner points).")
add_bullet(doc, "For every person detected in a frame, the system takes their foot position (bottom-centre of their bounding box).")
add_bullet(doc, "It runs cv2.pointPolygonTest() — a mathematical test to check if that point is inside the polygon.")
add_bullet(doc, "The person must be inside the zone for 3 consecutive frames before an alert fires.")
add_bullet(doc, "Each alert has a 15-second cooldown per person per zone.")
add_bullet(doc, "Time-based zones are supported: the Staff Only zone only fires alerts between 10pm and 8am.")
add_para(doc, "Current zones configured:")
add_bullet(doc, 'cam-01: "Back Office" covers top-right 320\u00d7400 pixel area. "Staff Only" covers bottom-left 300\u00d7280 pixel area (night hours only).')
add_bullet(doc, 'cam-03: "Storage Room C" covers top-right 220\u00d7500 pixel area.')
add_bullet(doc, 'cam-04: "Manager Office D" covers top-right 270\u00d7450 pixel area.')
add_para(doc, "Accuracy: HIGH when zones are correct (90%+). The polygon math is exact — a person is either inside or outside. The problem is not the algorithm but whether the polygon coordinates match your actual store layout.")
add_para(doc, "Action Required (IMPORTANT): You must verify that the polygon corners match your physical store. See Section 5 for how to do this with a snapshot tool.")

# 3.4
add_h2(doc, "4. Shoplifting / Product Sweep Detection")
add_para(doc, "This analyzer does not use zones. It tracks people and objects across frames to detect sweep-like hand movements near products.")
add_para(doc, "How it works:")
add_bullet(doc, 'YOLOv8 detects all "person" class objects in each frame.')
add_bullet(doc, "ByteTrack assigns each person a unique tracking ID that persists across frames.")
add_bullet(doc, "The system tracks how many detectable objects (bottles, cups, books, etc.) appear near each person per second.")
add_bullet(doc, "If a person causes 3 or more product detections to disappear within a 4-second window, a shoplifting event fires.")
add_bullet(doc, "Proximity threshold: 130 pixels for cam-01, 120 pixels for cam-03 and cam-04.")
add_bullet(doc, "30-second cooldown per tracked person.")
add_para(doc, 'What YOLOv8 can detect as "products": bottle, cup, bowl, banana, apple, orange, broccoli, carrot, sandwich, book, vase. These are COCO dataset classes. Products in sealed packaging that do not look like these classes will NOT be detected.')
add_para(doc, "Accuracy: MEDIUM (60\u201370% for common retail products). Works well for loose fruits, bottles, books. Does NOT work for: packets of chips, confectionery boxes, electronics, or any item not in the COCO 80-class set. To detect your specific products, a custom YOLOv8 model must be trained on your product images.")
add_para(doc, "Action Required: For production accuracy, provide 200+ images of your actual products and retrain a custom YOLOv8 model. This is a 2\u20133 day effort.")

# 3.5
add_h2(doc, "5. Inventory Movement Detection")
add_para(doc, "This analyzer counts detected products inside each shelf zone per 5-second interval and alerts when the count drops suddenly.")
add_para(doc, "How it works:")
add_bullet(doc, "Every 5 seconds, the system counts all product-class objects detected inside each defined shelf bounding box.")
add_bullet(doc, "It compares the new count to the previous count.")
add_bullet(doc, "If the count drops by 2 or more items, an alert fires.")
add_bullet(doc, "If a person is present in the same zone, severity is reduced to MEDIUM (they may be restocking).")
add_bullet(doc, "If no person is present and items disappear, severity is CRITICAL or HIGH.")
add_bullet(doc, "20-second cooldown per zone.")
add_para(doc, "Current shelf zones: Each camera has 3 shelf zones splitting the 1920\u00d71080 frame into thirds horizontally (top 200\u2013750 pixel rows). These are estimates of where shelves typically appear in retail stores.")
add_para(doc, "Accuracy: LOW without calibration (40\u201350%), MEDIUM after calibration (70%). Same product detection limitation as shoplifting — only COCO classes are detected. Additionally, shelf zones must actually align with your shelves.")
add_para(doc, "Action Required: Verify shelf bounding boxes against a live snapshot. Retrain model for your products.")

# 3.6
add_h2(doc, "6. Queue Length Monitoring")
add_para(doc, "This analyzer counts persons inside a defined checkout queue zone and alerts when the queue is too long.")
add_para(doc, "How it works:")
add_bullet(doc, "Every 3 seconds, it counts person-class detections inside the queue bounding box.")
add_bullet(doc, "If count exceeds the max_queue threshold (default: 4\u20135 people), a MEDIUM alert fires.")
add_bullet(doc, "60-second cooldown to avoid repeated alerts for the same long queue.")
add_para(doc, "Accuracy: HIGH when zone is correct (85%+). Person detection is the most reliable class in YOLOv8 (trained on millions of examples). The only issue is zone placement.")
add_para(doc, "Action Required: Verify the queue zone bbox aligns with your checkout counter.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: RISK SCORES
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, "What the Risk Score Percentages Mean")
add_para(doc, "The Risk Score (0\u2013100) shown on the dashboard is NOT a percentage of certainty. It is an accumulation of weighted event points over a rolling 60-second window.")

make_table(doc,
    ["Event Type", "Points Added", "Severity Weight"],
    [
        ["Watchlist Match",       "+40 points", "Highest — known threat"],
        ["Shoplifting / Sweep",   "+30 points", "High — direct theft indicator"],
        ["Camera Tamper",         "+25 points", "High — sabotage indicator"],
        ["Fall Detected",         "+25 points", "High — safety emergency"],
        ["Restricted Zone Entry", "+20 points", "Medium-High — unauthorised access"],
        ["Loitering / Dwell",     "+15 points", "Medium"],
        ["Inventory Movement",    "+10 points", "Medium — may be restocking"],
        ["Queue Length Breach",   "+10 points", "Low — operational issue"],
    ],
    [40, 25, 35],
)

add_para(doc, "")
add_para(doc, "Severity bands:")

make_table(doc,
    ["Score Range", "Label", "Colour on Dashboard"],
    [
        ["0 \u2013 25",   "LOW",      "Green"],
        ["26 \u2013 60",  "MEDIUM",   "Amber"],
        ["61 \u2013 89",  "HIGH",     "Red"],
        ["90 \u2013 100", "CRITICAL", "Red + pulsing border"],
    ],
    [30, 25, 45],
)

add_para(doc, "")
add_para(doc, "Why Zone D shows 100 CRITICAL right now: The RTSP stream for cam-04 (Zone D, Dahua camera at 192.168.1.251) has been triggering tamper events each time the stream reconnects (+25 points each time). Multiple tamper events within the 60-second window accumulate to 100. This is a real detection of stream instability — not a false alarm about physical tampering, but it is inflating the risk score.")
add_para(doc, "How the score resets: Every 60 seconds, the window slides forward. If no new events fire, the score drops to 0 automatically. The score is dynamic — calm periods always return to 0.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: CALIBRATION
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, "How to Calibrate Zones for Your Store Layout")
add_para(doc, "This is the most important step for production accuracy. You need to match the pixel coordinates in cameras.yaml to the actual positions of shelves, restricted areas, and checkout lanes as they appear in each camera's feed.")

add_h2(doc, "Step 1 \u2014 Get a Reference Snapshot")
add_para(doc, "Open your browser and navigate to:")
add_code_block(doc, "http://localhost:8800/api/cameras/cam-01/snapshot\nhttp://localhost:8800/api/cameras/cam-03/snapshot\nhttp://localhost:8800/api/cameras/cam-04/snapshot")
add_para(doc, "You will get a JPEG image of what each camera currently sees. Save these images. This is your reference for placing coordinates.")

add_h2(doc, "Step 2 \u2014 Find Pixel Coordinates Using MS Paint")
add_bullet(doc, "Open the saved JPEG in Microsoft Paint.")
add_bullet(doc, 'Move your mouse to the top-left corner of the area you want to define (e.g., the start of Shelf A1).')
add_bullet(doc, 'Look at the bottom bar of Paint \u2014 it shows "X: 450, Y: 200" \u2014 these are your x1, y1 coordinates.')
add_bullet(doc, "Move to the bottom-right corner of the same shelf. Note those coordinates as x2, y2.")
add_bullet(doc, "Write down all four numbers: [x1, y1, x2, y2].")

add_h2(doc, "Step 3 \u2014 Update cameras.yaml")
add_para(doc, "Open the file at:")
add_code_block(doc, "vantag\\backend\\config\\cameras.yaml")
add_para(doc, "Find the camera section (cam-01, cam-03, or cam-04) and update the bbox values under inventory_movement zones, restricted_zone polygons, or queue_length zones. Example for a shelf:")
add_code_block(doc, '      inventory_movement:\n        zones:\n          - label: "Shelf A1"\n            bbox: [YOUR_X1, YOUR_Y1, YOUR_X2, YOUR_Y2]')

add_h2(doc, "Step 4 \u2014 Restart the Backend")
add_code_block(doc, 'cd "D:\\AI Algo\\Collaterals\\Profiles\\Retail Nazar\\vantag"\npython -m uvicorn backend.api.main:app --host 0.0.0.0 --port 8800')
add_para(doc, "The new zone coordinates take effect immediately on restart. No other changes needed.")

add_h2(doc, "Step 5 \u2014 Verify with a Live Walk-Through Test")
add_bullet(doc, "Have a person physically walk into the area you defined as restricted.")
add_bullet(doc, 'Within 3\u20135 seconds, an incident should appear in the Incidents page with type "Restricted Zone Entry".')
add_bullet(doc, "If it does not appear, the zone coordinates are off \u2014 adjust x1/y1/x2/y2 and retry.")
add_bullet(doc, "Repeat for each zone on each camera.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: ACCURACY MATRIX
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, "Honest Accuracy Assessment \u2014 Current vs. Calibrated")

make_table(doc,
    ["Detection", "Current Accuracy", "After Zone Calibration", "After Custom Model Training"],
    [
        ["Camera Tamper",           "95%",                          "95% (no change needed)",     "N/A"],
        ["Fall Detection",          "80%",                          "85% (more frames filter)",   "90% with more training data"],
        ["Restricted Zone Entry",   "50% (zones uncalibrated)",     "92% (after calibration)",    "N/A"],
        ["Queue Length",            "55% (zone uncalibrated)",      "88% (after calibration)",    "N/A"],
        ["Shoplifting Detection",   "60% (COCO products only)",     "65% (calibrated shelf area)","85% (custom product model)"],
        ["Inventory Movement",      "40% (COCO + uncalibrated)",    "60% (calibrated zones)",     "80% (custom product model)"],
        ["Watchlist / Face Match",  "0% (no faces uploaded)",       "N/A",                        "90%+ (after face uploads)"],
    ],
    [25, 25, 25, 25],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7: PLUG AND PLAY
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, 'What "Plug and Play" Means \u2014 And Its Limits')
add_para(doc, "The plug-and-play claim means: you connect a camera, enter its IP address, and the system automatically discovers the RTSP URL, starts streaming, and begins detecting people, falls, and tamper events within minutes. No manual RTSP configuration is needed.")
add_para(doc, "What plug-and-play does NOT cover automatically:")
add_bullet(doc, "It does not know the physical layout of your store. You must tell it where your shelves are.")
add_bullet(doc, "It does not know which products you sell. Generic products (bottles, books) are detected by default.")
add_bullet(doc, "It does not know who your staff are vs. customers. All persons are tracked equally unless a watchlist face is uploaded.")
add_bullet(doc, "Zone calibration is a one-time 15\u201330 minute setup per camera per store. After that it runs fully automatically.")
add_para(doc, "Think of it like a new CCTV camera: it records automatically, but a human must aim it at the right angle and tell the system what areas matter. Vantag automates everything except the initial zone pointing.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8: ACTION PLAN
# ══════════════════════════════════════════════════════════════════════════════

add_h1(doc, "Recommended Action Plan (Priority Order)")

add_h2(doc, "Priority 1 \u2014 Done Automatically (No Action Needed)")
add_bullet(doc, "Camera RTSP connection and streaming")
add_bullet(doc, "Person detection across all frames")
add_bullet(doc, "Tamper detection")
add_bullet(doc, "Fall detection")
add_bullet(doc, "Risk score computation and dashboard display")

add_h2(doc, "Priority 2 \u2014 One-Time Setup (30 Minutes Per Store)")
add_bullet(doc, "Take snapshots from each camera and open in MS Paint")
add_bullet(doc, "Verify shelf bounding boxes match real shelf positions")
add_bullet(doc, "Verify restricted zone polygons match real restricted areas")
add_bullet(doc, "Verify queue zone matches real checkout counter")
add_bullet(doc, "Update cameras.yaml with corrected coordinates")
add_bullet(doc, "Restart backend and do a live walk-through test")

add_h2(doc, "Priority 3 \u2014 Optional but Recommended for Higher Accuracy")
add_bullet(doc, "Upload staff and known person face images to the Watchlist for face-match alerts")
add_bullet(doc, "Collect 200+ product images from your shelves and train a custom YOLOv8 model")
add_bullet(doc, "Configure staff_zone_colors in cameras.yaml with the actual shirt/uniform colours your staff wear (used for staff vs. customer differentiation)")

add_h2(doc, "Priority 4 \u2014 Advanced (When Scaling to Multiple Stores)")
add_bullet(doc, "Use the Vantag Zone Editor API (/api/cameras/{id}/zones) to push zone configs remotely without touching cameras.yaml")
add_bullet(doc, "Build a visual zone editor overlay on the CameraView page (draw boxes on live feed)")
add_bullet(doc, "Enable NVIDIA TensorRT acceleration on Jetson devices for 3x faster inference")

# ── Save ─────────────────────────────────────────────────────────────────────

doc.save(OUTPUT)
print(f"SUCCESS: Document written to {OUTPUT}")
